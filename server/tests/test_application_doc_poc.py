"""申请表模板 POC 测试：验证参考模板可被 python-docx 读取与写回。"""
import tempfile
import unittest
from pathlib import Path


class TestApplicationDocPOC(unittest.TestCase):
    def setUp(self):
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            self.skipTest("python-docx 未安装，跳过 POC 测试")

        self.template_path = Path(__file__).resolve().parents[2] / "template" / "基于SpringBoot的鱼缸控制系统_申请表.docx"
        if not self.template_path.exists():
            self.skipTest(f"参考模板不存在: {self.template_path}")

    def test_template_can_be_opened(self):
        from docx import Document

        doc = Document(str(self.template_path))
        self.assertGreater(len(doc.paragraphs), 0, "模板段落不应为空")
        self.assertGreater(len(doc.tables), 0, "申请表模板应包含表格")

    def test_template_can_be_saved_after_edit(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "申请表_poc_写回验证.docx"
            doc = Document(str(self.template_path))
            doc.paragraphs[0].add_run(" [POC写回验证]")
            doc.save(str(output))

            reopened = Document(str(output))
            first_text = reopened.paragraphs[0].text
            self.assertIn("POC写回验证", first_text)


if __name__ == "__main__":
    unittest.main()
