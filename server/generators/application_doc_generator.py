"""申请表生成器 - MVP版"""
from pathlib import Path

from config import BASE_DIR, Config
from generators.models import ProjectContext


class ApplicationDocGenerator:
    MANUAL_FIELDS = ["著作权人", "证件号码", "联系地址", "联系电话", "权利取得方式", "权利范围"]
    TEMPLATE_CANDIDATES = [
        Config.TEMPLATE_DIR / "application_template.docx",
        BASE_DIR.parent / "template" / "基于SpringBoot的鱼缸控制系统_申请表.docx",
    ]

    def generate(self, task_id: str, context: ProjectContext) -> str:
        try:
            from docx import Document
            from docx.enum.text import WD_COLOR_INDEX
        except ImportError:
            return self._fallback_text(task_id, context)

        template_path = self._resolve_template_path()
        doc = Document(template_path) if template_path else Document()

        mapping = {
            "{{SOFTWARE_NAME}}": context.software_name,
            "{{SHORT_NAME}}": context.short_name,
            "{{VERSION}}": "V1.0",
            "{{COMPLETION_DATE}}": context.completion_date,
            "{{TOTAL_LINES}}": str(context.total_lines),
            "{{LANGUAGES}}": context.tech_config.get("languages", ""),
            "{{FEATURE_SUMMARY}}": context.feature_summary,
            "{{OS}}": context.tech_config.get("os", ""),
            "{{RUNTIME}}": context.tech_config.get("runtime", ""),
            "{{DEV_TOOLS}}": context.tech_config.get("dev_tools", ""),
            "{{MANUAL_FIELDS}}": "、".join(self.MANUAL_FIELDS),
        }
        self._replace_placeholders(doc, mapping)
        self._highlight_manual_fields(doc, WD_COLOR_INDEX.YELLOW)

        out = Config.OUTPUT_DIR / task_id
        out.mkdir(parents=True, exist_ok=True)
        path = out / f"{self._safe(context.software_name)}_申请表.docx"
        doc.save(path)
        return str(path)

    def _fallback_text(self, task_id: str, context: ProjectContext) -> str:
        out = Config.OUTPUT_DIR / task_id
        out.mkdir(parents=True, exist_ok=True)
        path = out / f"{self._safe(context.software_name)}_申请表.txt"
        text = [
            "计算机软件著作权登记申请表（自动生成草稿）\n\n",
            f"软件全称：{context.software_name}\n",
            f"软件简称：{context.short_name}\n",
            "版本号：V1.0\n",
            f"开发完成日期：{context.completion_date}\n",
            f"源程序量：{context.total_lines}\n",
            f"编程语言：{context.tech_config.get('languages', '')}\n",
            f"主要功能：{context.feature_summary}\n\n",
            "以下字段需手动填写：\n",
        ]
        text.extend([f"- {f}\n" for f in self.MANUAL_FIELDS])
        path.write_text("".join(text), encoding="utf-8")
        return str(path)

    def _resolve_template_path(self) -> Path | None:
        for candidate in self.TEMPLATE_CANDIDATES:
            if candidate.exists():
                return candidate
        return None

    def _replace_placeholders(self, doc, mapping: dict[str, str]):
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, mapping)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, mapping)

    def _replace_in_paragraph(self, paragraph, mapping: dict[str, str]):
        text = paragraph.text
        replaced = text
        for key, value in mapping.items():
            replaced = replaced.replace(key, str(value))
        if replaced == text:
            return
        paragraph.clear()
        paragraph.add_run(replaced)

    def _highlight_manual_fields(self, doc, color):
        for paragraph in doc.paragraphs:
            self._highlight_keywords(paragraph, color)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._highlight_keywords(paragraph, color)

    def _highlight_keywords(self, paragraph, color):
        text = paragraph.text
        if not text:
            return
        paragraph.clear()
        for field in self.MANUAL_FIELDS:
            text = text.replace(field, f"[[[{field}]]]")
        segments = text.split("[[[")
        first = True
        for segment in segments:
            if first:
                paragraph.add_run(segment)
                first = False
                continue
            if "]]]" not in segment:
                paragraph.add_run("[[[" + segment)
                continue
            keyword, tail = segment.split("]]]", 1)
            run = paragraph.add_run(keyword)
            run.font.highlight_color = color
            paragraph.add_run(tail)

    def _safe(self, name: str) -> str:
        return "".join(ch if ch not in '\\/:*?\"<>|' else "_" for ch in name).strip() or "软著材料"
