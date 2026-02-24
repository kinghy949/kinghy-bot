"""源码文档生成器"""
import re

from config import Config
from generators.docx_utils import apply_standard_header
from generators.models import ProjectContext


class SourceDocGenerator:
    MAX_PAGES = 70
    LINES_PER_PAGE = 50
    FULL_DUMP_THRESHOLD = 3000
    HEADER_FILE_COUNT = 2
    TAIL_FILE_COUNT = 2

    def generate(self, task_id: str, context: ProjectContext) -> str:
        try:
            from docx import Document
            from docx.enum.text import WD_BREAK
            from docx.shared import Cm, Pt
        except ImportError:
            return self._fallback_text(task_id, context)

        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

        apply_standard_header(doc, context.software_name, context.software_version)
        doc.add_heading(f"{context.software_name} 源程序文档", level=1)

        all_files = self._build_file_blocks(context)
        selected_files, strategy = self._select_files(context, all_files)
        page_line_count = 0
        for item in selected_files:
            line_no = 1
            link_mark = self._feature_link_mark(context, item["path"])
            heading = item["path"] if not link_mark else f"{item['path']}（关联功能点：{link_mark}）"
            doc.add_heading(heading, level=2)
            for line in item["lines"]:
                if page_line_count == self.LINES_PER_PAGE:
                    page_break_p = doc.add_paragraph("")
                    page_break_p.add_run().add_break(WD_BREAK.PAGE)
                    page_line_count = 0
                p = doc.add_paragraph(f"{line_no:04d}  {line}")
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10.5)
                line_no += 1
                page_line_count += 1

        selected_line_count = sum(item["line_count"] for item in selected_files)
        feature_to_files, file_to_features = self._build_traceability(context, selected_files)
        context.doc_metrics["source"] = {
            "strategy": strategy,
            "total_code_lines": sum(item["line_count"] for item in all_files),
            "selected_code_lines": selected_line_count,
            "selected_files": [item["path"] for item in selected_files],
            "page_limit": self.MAX_PAGES,
            "lines_per_page": self.LINES_PER_PAGE,
            "estimated_pages": (selected_line_count + self.LINES_PER_PAGE - 1) // self.LINES_PER_PAGE,
            "last_file_complete": True,
            "traceability": {
                "feature_to_files": feature_to_files,
                "file_to_features": file_to_features,
            },
            "copyright": self._extract_copyright_notice(all_files),
        }

        out = Config.OUTPUT_DIR / task_id
        out.mkdir(parents=True, exist_ok=True)
        path = out / f"{self._safe(context.software_name)}_源程序.docx"
        doc.save(path)
        return str(path)

    def _build_file_blocks(self, context: ProjectContext) -> list[dict]:
        blocks: list[dict] = []
        for path in sorted(context.generated_code.keys()):
            lines = context.generated_code[path].splitlines()
            blocks.append({"path": path, "lines": lines, "line_count": len(lines)})
        return blocks

    def _select_files(self, context: ProjectContext, blocks: list[dict]) -> tuple[list[dict], str]:
        total_lines = sum(item["line_count"] for item in blocks)
        if total_lines <= self.FULL_DUMP_THRESHOLD:
            return blocks, "full"

        max_lines = self.MAX_PAGES * self.LINES_PER_PAGE
        selected: list[dict] = []
        selected_paths: set[str] = set()
        selected_lines = 0

        def try_add(item: dict):
            nonlocal selected_lines
            if item["path"] in selected_paths:
                return
            if selected_lines + item["line_count"] > max_lines:
                return
            selected.append(item)
            selected_paths.add(item["path"])
            selected_lines += item["line_count"]

        for item in blocks[: self.HEADER_FILE_COUNT]:
            try_add(item)
        for item in blocks[-self.TAIL_FILE_COUNT :]:
            try_add(item)

        for item in self._sort_core_blocks(context, blocks):
            try_add(item)

        for item in blocks:
            try_add(item)

        if not selected:
            selected = blocks[:1]
        selected.sort(key=lambda x: x["path"])
        return selected, "segment"

    def _sort_core_blocks(self, context: ProjectContext, blocks: list[dict]) -> list[dict]:
        keywords = ("api", "service", "controller", "view", "module", "model", "store", "router")
        feature_terms = [self._slug(f.name) for f in context.feature_list]
        scored: list[tuple[int, str, dict]] = []
        for item in blocks:
            path_lc = item["path"].lower()
            score = sum(2 for kw in keywords if kw in path_lc)
            for term in feature_terms:
                if term and term in self._slug(path_lc):
                    score += 3
            scored.append((score, item["path"], item))
        scored.sort(key=lambda x: (-x[0], x[1]))
        return [item for score, _, item in scored if score > 0]

    def _feature_link_mark(self, context: ProjectContext, file_path: str) -> str:
        marks: list[str] = []
        slugged_path = self._slug(file_path)
        for idx, feature in enumerate(context.feature_list, start=1):
            fid = feature.feature_id or f"F{idx:02d}"
            if file_path in feature.code_files:
                marks.append(fid)
                continue
            feature_slug = self._slug(feature.name)
            if feature_slug and feature_slug in slugged_path:
                marks.append(fid)
                continue
            if feature.name and feature.name in file_path:
                marks.append(fid)
        return ",".join(dict.fromkeys(marks))

    def _build_traceability(self, context: ProjectContext, selected_files: list[dict]) -> tuple[dict, dict]:
        selected_paths = [item["path"] for item in selected_files]
        feature_to_files: dict[str, list[str]] = {}
        file_to_features: dict[str, list[str]] = {}
        for idx, feature in enumerate(context.feature_list, start=1):
            fid = feature.feature_id or f"F{idx:02d}"
            linked = [path for path in selected_paths if path in feature.code_files]
            if not linked:
                linked = [path for path in selected_paths if self._slug(feature.name) in self._slug(path)]
            feature_to_files[fid] = linked
            for path in linked:
                file_to_features.setdefault(path, [])
                if fid not in file_to_features[path]:
                    file_to_features[path].append(fid)
        return feature_to_files, file_to_features

    def _slug(self, text: str) -> str:
        cleaned = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff]+", "_", text).strip("_")
        return cleaned.lower()

    def _fallback_text(self, task_id: str, context: ProjectContext) -> str:
        out = Config.OUTPUT_DIR / task_id
        out.mkdir(parents=True, exist_ok=True)
        path = out / f"{self._safe(context.software_name)}_源程序.txt"
        all_files = self._build_file_blocks(context)
        selected, strategy = self._select_files(context, all_files)
        chunks = [f"{context.software_name} {context.software_version} 源程序文档\n"]
        chunks.append(f"输出策略：{strategy}\n")
        for item in selected:
            chunks.append(f"\n## {item['path']}\n")
            for i, line in enumerate(item["lines"], start=1):
                chunks.append(f"{i:04d} {line}\n")
        feature_to_files, file_to_features = self._build_traceability(context, selected)
        context.doc_metrics["source"] = {
            "strategy": strategy,
            "total_code_lines": sum(item["line_count"] for item in all_files),
            "selected_code_lines": sum(item["line_count"] for item in selected),
            "selected_files": [item["path"] for item in selected],
            "page_limit": self.MAX_PAGES,
            "lines_per_page": self.LINES_PER_PAGE,
            "estimated_pages": (sum(item["line_count"] for item in selected) + self.LINES_PER_PAGE - 1) // self.LINES_PER_PAGE,
            "last_file_complete": True,
            "traceability": {
                "feature_to_files": feature_to_files,
                "file_to_features": file_to_features,
            },
            "copyright": self._extract_copyright_notice(all_files),
        }
        path.write_text("".join(chunks), encoding="utf-8")
        return str(path)

    def _safe(self, name: str) -> str:
        return "".join(ch if ch not in '\\/:*?\"<>|' else "_" for ch in name).strip() or "软著材料"

    def _extract_copyright_notice(self, blocks: list[dict]) -> dict:
        owners: set[str] = set()
        patterns = [
            re.compile(r"copyright\s*(?:\(c\))?\s*[:：]?\s*(.+)", re.IGNORECASE),
            re.compile(r"版权所有\s*[:：]?\s*(.+)"),
            re.compile(r"著作权(?:归属|所有者)?\s*[:：]?\s*(.+)"),
        ]
        for item in blocks:
            for line in item["lines"][:80]:
                text = line.strip()
                if not text:
                    continue
                for pattern in patterns:
                    matched = pattern.search(text)
                    if not matched:
                        continue
                    owner = matched.group(1).strip().strip("*/#- ")
                    if owner:
                        owners.add(owner[:120])
                    break
        return {
            "has_notice": bool(owners),
            "owners": sorted(owners),
        }
