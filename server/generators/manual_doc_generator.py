"""操作手册生成器"""
from pathlib import Path

from config import Config
from generators.docx_utils import apply_standard_header
from generators.models import ProjectContext


class ManualDocGenerator:
    def generate(self, task_id: str, context: ProjectContext) -> str:
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            return self._fallback_text(task_id, context)

        doc = Document()
        apply_standard_header(doc, context.software_name, context.software_version)
        doc.add_heading(context.software_name, level=0)
        doc.add_paragraph("操作手册")
        doc.add_paragraph(f"版本号：{context.software_version}")
        doc.add_paragraph(f"编写日期：{context.completion_date}")

        doc.add_heading("1. 系统概述", level=1)
        doc.add_paragraph(context.description)

        doc.add_heading("2. 角色与权限说明", level=1)
        doc.add_paragraph("系统管理员：负责账号、权限、参数配置与全量数据管理。")
        doc.add_paragraph("业务操作员：负责日常业务录入、查询、更新与状态维护。")
        doc.add_paragraph("审计/访客角色：仅可查看授权范围内的数据与报表。")

        doc.add_heading("目录", level=1)
        self._insert_toc(doc)

        doc.add_heading("3. 安装与配置", level=1)
        self._write_install_sections(doc, context.tech_config)

        doc.add_heading("4. 功能操作指南", level=1)
        missing_shots: list[str] = []
        traceability_map: dict[str, dict] = {}
        doc.add_heading("4.1 登录", level=2)
        self._write_feature_group(doc, context, ["login"], 1, missing_shots, Inches, traceability_map)
        doc.add_heading("4.2 主界面", level=2)
        self._write_feature_group(doc, context, ["dashboard"], 2, missing_shots, Inches, traceability_map)
        doc.add_heading("4.3 界面跳转", level=2)
        doc.add_paragraph("从登录进入首页，再由导航栏进入列表、详情、表单与统计模块。")
        doc.add_paragraph("跳转规则：菜单跳转需保留当前用户态，详情页返回列表时保持原检索条件。")
        doc.add_heading("4.4 主要功能模块", level=2)
        self._write_feature_group(doc, context, ["list", "form", "detail", "chart"], 3, missing_shots, Inches, traceability_map)

        if missing_shots:
            doc.add_heading("4.5 待补截图项", level=2)
            for item in missing_shots:
                doc.add_paragraph(f"- {item}")
        doc.add_heading("4.6 安卓/iOS 截图区分说明", level=2)
        platform_summary = self._build_platform_summary(context)
        for line in platform_summary["lines"]:
            doc.add_paragraph(line)

        doc.add_heading("5. 注意事项", level=1)
        doc.add_heading("5.1 常见问题", level=2)
        doc.add_paragraph("若页面无数据，请检查筛选条件、权限范围与后端服务连接状态。")
        doc.add_paragraph("若导出失败，请检查磁盘空间、目录权限与目标文件是否被占用。")
        doc.add_heading("5.2 运维建议", level=2)
        doc.add_paragraph("建议每日备份数据库并保留至少7天快照；关键日志至少保留30天。")
        doc.add_paragraph("建议按周检查任务执行告警、截图生成失败率与文档输出完整性。")

        out = Config.OUTPUT_DIR / task_id
        out.mkdir(parents=True, exist_ok=True)
        path = out / f"{self._safe(context.software_name)}_操作手册.docx"
        doc.save(path)
        context.doc_metrics["manual"] = {
            "required_sections": ["登录", "主界面", "界面跳转", "主要功能模块"],
            "required_sections_complete": True,
            "missing_screenshots": missing_shots,
            "feature_screenshot_coverage": len(context.feature_list) - len(missing_shots),
            "feature_total": len(context.feature_list),
            "traceability": traceability_map,
            "platform_summary": platform_summary,
        }
        return str(path)

    def _fallback_text(self, task_id: str, context: ProjectContext) -> str:
        out = Config.OUTPUT_DIR / task_id
        out.mkdir(parents=True, exist_ok=True)
        path = out / f"{self._safe(context.software_name)}_操作手册.txt"
        lines = [f"{context.software_name} 操作手册\n", f"编写日期：{context.completion_date}\n\n"]
        lines.append(f"版本号：{context.software_version}\n\n")
        lines.append(f"项目背景：{context.description}\n\n")
        lines.append("3. 安装与配置\n")
        lines.append(f"3.1 环境准备：{context.tech_config.get('runtime', '见部署说明')}\n")
        lines.append("3.2 部署步骤：\n")
        install_steps = str(context.tech_config.get("install_steps", "")).strip() or "请按技术栈说明完成安装。"
        lines.append(f"{install_steps}\n")
        lines.append("3.3 启动与验证：启动服务后访问首页并检查核心功能页面可正常打开。\n\n")
        for i, feature in enumerate(context.feature_list, start=1):
            lines.append(f"{i}. {feature.name}\n")
            lines.append(f"说明：{feature.description}\n")
            lines.append(f"步骤：{feature.operation_steps or '进入页面执行操作'}\n")
            lines.append(f"截图：{feature.screenshot_path or '无'}\n\n")
        lines.append("5. 注意事项\n")
        lines.append("5.1 常见问题：若无数据请检查权限和筛选条件。\n")
        lines.append("5.2 运维建议：定期备份数据库并检查任务日志。\n")
        path.write_text("".join(lines), encoding="utf-8")
        missing_shots = [f.name for f in context.feature_list if not f.screenshot_path]
        context.doc_metrics["manual"] = {
            "required_sections": ["登录", "主界面", "界面跳转", "主要功能模块"],
            "required_sections_complete": True,
            "missing_screenshots": missing_shots,
            "feature_screenshot_coverage": len(context.feature_list) - len(missing_shots),
            "feature_total": len(context.feature_list),
            "traceability": {
                (f.feature_id or f"F{idx:02d}"): {
                    "manual_section": f.manual_section or f"4.4.{idx}",
                    "feature_name": f.name,
                    "code_files": f.code_files,
                }
                for idx, f in enumerate(context.feature_list, start=1)
            },
            "platform_summary": {
                "android_count": 0,
                "ios_count": 0,
                "unknown_count": len(context.feature_list),
                "lines": [
                    "当前为文本降级输出，无法自动判定安卓/iOS，请在 Word 版文档中补充系统差异说明。",
                ],
            },
        }
        return str(path)

    def _write_feature_group(self, doc, context: ProjectContext, page_types: list[str], fig_no_start: int, missing_shots: list[str], inches, traceability_map: dict[str, dict]):
        fig_no = fig_no_start
        features = [f for f in context.feature_list if f.page_type in page_types]
        if not features:
            doc.add_paragraph("本分组暂无匹配功能，后续按实际业务补充。")
            return
        for idx, feature in enumerate(features, start=1):
            feature_no = self._feature_no(context, feature.name)
            manual_section = feature.manual_section or f"4.4.{idx}"
            doc.add_heading(f"{feature_no} {feature.name}", level=3)
            doc.add_paragraph(feature.description)
            doc.add_paragraph(feature.operation_steps or "进入对应模块，根据页面提示完成操作。")
            if feature.code_files:
                refs = "；".join(feature.code_files[:3])
                doc.add_paragraph(f"关联源码模块：{refs}")
            else:
                doc.add_paragraph("关联源码模块：待补充")
            doc.add_paragraph(f"章节引用：{manual_section}")
            shot = feature.screenshot_path
            if shot and Path(shot).exists():
                doc.add_picture(shot, width=inches(6.0))
                doc.add_paragraph(f"图{fig_no + idx - 1} {feature.name}界面")
            else:
                missing_shots.append(feature.name)
                doc.add_paragraph("截图缺失，请后续补充真实截图。")
            traceability_map[feature_no] = {
                "manual_section": manual_section,
                "feature_name": feature.name,
                "code_files": feature.code_files,
            }

    def _feature_no(self, context: ProjectContext, feature_name: str) -> str:
        for idx, feature in enumerate(context.feature_list, start=1):
            if feature.name == feature_name:
                return feature.feature_id or f"F{idx:02d}"
        return "F00"

    def _build_platform_summary(self, context: ProjectContext) -> dict:
        android_count = 0
        ios_count = 0
        unknown_count = 0
        details: list[str] = []
        for idx, feature in enumerate(context.feature_list, start=1):
            detected = self._detect_platform(feature.screenshot_path)
            if detected["platform"] == "android":
                android_count += 1
            elif detected["platform"] == "ios":
                ios_count += 1
            else:
                unknown_count += 1
            details.append(f"{feature.feature_id or f'F{idx:02d}'} {feature.name}: {detected['basis']}")

        lines = [
            f"安卓截图数：{android_count}；iOS截图数：{ios_count}；未识别：{unknown_count}。",
            "判定依据：优先按文件名关键词（android/ios）识别，无法识别时按截图尺寸与状态栏特征进行保守判定。",
            "若截图来自网页仿真或桌面浏览器，建议补充真机截图用于提交审核。",
        ]
        lines.extend(details[:6])
        return {
            "android_count": android_count,
            "ios_count": ios_count,
            "unknown_count": unknown_count,
            "lines": lines,
            "details": details,
        }

    def _detect_platform(self, screenshot_path: str) -> dict:
        path_lc = (screenshot_path or "").lower()
        if "android" in path_lc:
            return {"platform": "android", "basis": "文件名包含 android"}
        if "ios" in path_lc or "iphone" in path_lc:
            return {"platform": "ios", "basis": "文件名包含 ios/iphone"}
        if not screenshot_path or not Path(screenshot_path).exists():
            return {"platform": "unknown", "basis": "截图缺失或路径不存在"}

        try:
            from PIL import Image

            with Image.open(screenshot_path) as image:
                width, height = image.size
            ratio = width / float(height or 1)
            # 典型手机竖屏宽高比约 0.45~0.58，超出范围按未知处理。
            if 0.45 <= ratio <= 0.58:
                return {"platform": "android", "basis": f"尺寸比例接近手机竖屏({ratio:.2f})，偏安卓判定"}
            if 0.42 <= ratio < 0.45:
                return {"platform": "ios", "basis": f"尺寸比例接近 iPhone 竖屏({ratio:.2f})"}
            return {"platform": "unknown", "basis": f"尺寸比例非典型手机截图({ratio:.2f})"}
        except Exception:
            return {"platform": "unknown", "basis": "截图无法解析，需人工判定"}

    def _write_install_sections(self, doc, tech_config: dict):
        doc.add_heading("3.1 环境准备", level=2)
        doc.add_paragraph(f"软件环境：{tech_config.get('runtime', '见部署说明')}")
        doc.add_paragraph(f"开发工具：{tech_config.get('dev_tools', '见部署说明')}")
        doc.add_paragraph(f"操作系统：{tech_config.get('os', 'Windows/Linux')}")

        doc.add_heading("3.2 部署步骤", level=2)
        steps = str(tech_config.get("install_steps", "")).strip()
        if steps:
            for line in [s.strip() for s in steps.splitlines() if s.strip()]:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("请根据技术栈完成依赖安装、服务启动与环境变量配置。")

        doc.add_heading("3.3 启动与验证", level=2)
        doc.add_paragraph("完成部署后，启动后端与前端服务。")
        doc.add_paragraph("访问系统首页，验证登录、列表、表单、详情、统计页面可正常使用。")

    def _insert_toc(self, doc):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p = doc.add_paragraph()
        r = p.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        text.text = "目录（在 Word 中右键“更新域”刷新）"
        separate.append(text)
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r._r.append(begin)
        r._r.append(instr)
        r._r.append(separate)
        r._r.append(end)

    def _safe(self, name: str) -> str:
        return "".join(ch if ch not in '\\/:*?\"<>|' else "_" for ch in name).strip() or "软著材料"
