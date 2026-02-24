"""Word 文档公共工具。"""


def add_word_field(paragraph, field_name: str):
    """插入 Word 域（PAGE/NUMPAGES 等）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def apply_standard_header(doc, software_name: str, software_version: str):
    """统一页眉：软件全称 + 版本号 + 页码。"""
    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            paragraph = header.add_paragraph()
        else:
            paragraph = header.paragraphs[0]
            paragraph.clear()
        paragraph.add_run(f"{software_name} {software_version} 第")
        add_word_field(paragraph, "PAGE")
        paragraph.add_run("页")
